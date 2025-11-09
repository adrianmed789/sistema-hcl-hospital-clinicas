import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox, filedialog
from tkcalendar import DateEntry
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
import os
from datetime import datetime, date
import re


class SistemaDocumentosMedicos:
    def __init__(self):
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        # Ventana principal
        self.root = ctk.CTk()
        self.root.title("Sistema de Gestión de Documentos Médicos")
        self.root.geometry("1000x800")

        # Variables de los campos
        self.nombres_var = tk.StringVar()
        self.apellido_materno_var = tk.StringVar()
        self.apellido_paterno_var = tk.StringVar()
        self.fecha_nacimiento_var = tk.StringVar()
        self.edad_var = tk.StringVar()
        self.genero_var = tk.StringVar(value="Masculino")
        self.num_historia_var = tk.StringVar()
        self.num_registro_var = tk.StringVar()
        self.ocupacion_var = tk.StringVar()
        self.estado_civil_var = tk.StringVar()
        self.residencia_var = tk.StringVar()
        self.domicilio_var = tk.StringVar()
        self.procedencia_var = tk.StringVar()
        self.n_referencia1_var = tk.StringVar()
        self.referencia1_var = tk.StringVar()
        self.n_referencia2_var = tk.StringVar()
        self.referencia2_var = tk.StringVar()
        self.diagnosticos_var = tk.StringVar()
        self.diag_recetas_labs_var = tk.StringVar()
        self.cie10_var = tk.StringVar()
        self.servicio_var = tk.StringVar(value="Hematología")

        # Directorios y archivos
        self.directorio_base = os.getcwd()
        self.plantillas_dir = os.path.join(self.directorio_base, "PLANTILLAS")
        self.excel_file = os.path.join(self.directorio_base, "pacientes.xlsx")

        # Inicialización
        self._crear_estructura_directorios()
        self._build_interface()
        self._configurar_eventos()

    def _crear_estructura_directorios(self):
        """Crea las carpetas de plantillas y el archivo Excel si no existen."""
        try:
            subdirs = [
                "Consulta", "Interconsulta", "Recetas", 
                "Laboratorios", "Procedimientos", "Informes"
            ]
            for sub in subdirs:
                os.makedirs(os.path.join(self.plantillas_dir, sub), exist_ok=True)

            if not os.path.exists(self.excel_file):
                df = pd.DataFrame(columns=[
                    'Fecha_Registro', 'Nombre_Completo', 'Num_Historia', 'Num_Registro',
                    'Edad', 'Servicio', 'Diagnostico', 'Fecha_Internacion', 'Ruta_Carpeta'
                ])
                df.to_excel(self.excel_file, index=False)
        except Exception as e:
            messagebox.showerror("Error", f"Error creando estructura: {e}")

    def _build_interface(self):
        """Construye toda la interfaz gráfica."""
        main_frame = ctk.CTkScrollableFrame(self.root)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Título
        ctk.CTkLabel(
            main_frame,
            text="SISTEMA DE GESTIÓN DE DOCUMENTOS MÉDICOS",
            font=ctk.CTkFont(size=24, weight="bold")
        ).pack(pady=(0, 20))

        # ─── DATOS DEL PACIENTE ─────────────────────────────────────────
        datos_frame = ctk.CTkFrame(main_frame)
        datos_frame.pack(fill="x", pady=(0, 10))
        ctk.CTkLabel(
            datos_frame,
            text="DATOS DEL PACIENTE",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=10)

        gd = ctk.CTkFrame(datos_frame)
        gd.pack(fill="x", padx=20, pady=(0, 20))

        # Nombres
        ctk.CTkLabel(gd, text="Nombres:").grid(row=0, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.nombres_var, width=250).grid(
            row=0, column=1, columnspan=3, sticky="ew", padx=10, pady=5
        )

        # Apellidos
        ctk.CTkLabel(gd, text="Apellido Paterno:").grid(row=1, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.apellido_paterno_var, width=200).grid(
            row=1, column=1, sticky="ew", padx=10, pady=5
        )
        ctk.CTkLabel(gd, text="Apellido Materno:").grid(row=1, column=2, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.apellido_materno_var, width=200).grid(
            row=1, column=3, sticky="ew", padx=10, pady=5
        )

        # Fecha de Nacimiento y Edad
        ctk.CTkLabel(gd, text="Fecha Nacimiento (DD/MM/YY):").grid(
            row=2, column=0, sticky="w", padx=10, pady=5
        )
        ctk.CTkEntry(
            gd,
            textvariable=self.fecha_nacimiento_var,
            width=120,
            placeholder_text="DD/MM/YY"
        ).grid(row=2, column=1, sticky="ew", padx=10, pady=5)
        ctk.CTkLabel(gd, text="Edad:").grid(row=2, column=2, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.edad_var, width=100, state="readonly").grid(
            row=2, column=3, sticky="ew", padx=10, pady=5
        )

        # Fecha de Internación
        ctk.CTkLabel(gd, text="Fecha Internación:").grid(
            row=3, column=0, sticky="w", padx=10, pady=5
        )
        self.fecha_internacion_entry = DateEntry(
            gd, width=12, background='darkblue', foreground='white',
            borderwidth=2, date_pattern='dd/mm/yyyy'
        )
        self.fecha_internacion_entry.grid(row=3, column=1, sticky="ew", padx=10, pady=5)

        # Número de Historia Clínica
        ctk.CTkLabel(gd, text="N° Historia Clínica:").grid(row=4, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.num_historia_var, width=150).grid(
            row=4, column=1, sticky="ew", padx=10, pady=5
        )
        
        # Número de Registro (calculado automáticamente)
        ctk.CTkLabel(gd, text="N° Registro:").grid(row=4, column=2, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.num_registro_var, width=150, state="readonly").grid(
            row=4, column=3, sticky="ew", padx=10, pady=5
        )

        # Ocupación y Estado Civil
        ctk.CTkLabel(gd, text="Ocupación:").grid(row=5, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.ocupacion_var, width=200).grid(
            row=5, column=1, sticky="ew", padx=10, pady=5
        )
        ctk.CTkLabel(gd, text="Estado Civil:").grid(row=5, column=2, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.estado_civil_var, width=150).grid(
            row=5, column=3, sticky="ew", padx=10, pady=5
        )

        # Referencias
        ctk.CTkLabel(gd, text="Referencia 1:").grid(row=6, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.referencia1_var, width=200).grid(
            row=6, column=1, sticky="ew", padx=10, pady=5
        )
        ctk.CTkLabel(gd, text="Referencia 2:").grid(row=6, column=2, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(gd, textvariable=self.referencia2_var, width=200).grid(
            row=6, column=3, sticky="ew", padx=10, pady=5
        )

        gd.grid_columnconfigure(1, weight=1)
        gd.grid_columnconfigure(3, weight=1)

        # ─── DATOS CLÍNICOS ────────────────────────────────────────────
        cf = ctk.CTkFrame(main_frame)
        cf.pack(fill="x", pady=(0, 10))
        ctk.CTkLabel(
            cf, text="DATOS CLÍNICOS", font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=10)

        cg = ctk.CTkFrame(cf)
        cg.pack(fill="x", padx=20, pady=(0, 20))

        # Servicio (Desplegable)
        ctk.CTkLabel(cg, text="Servicio:").grid(row=0, column=0, sticky="w", padx=10, pady=5)
        servicios = ["Hematología", "Medicina Interna", "Oncología Clínica", "Oncología Quirúrgica"]
        ctk.CTkComboBox(
            cg, 
            values=servicios,
            variable=self.servicio_var,
            width=200
        ).grid(row=0, column=1, sticky="ew", padx=10, pady=5)

        # Diagnósticos
        ctk.CTkLabel(cg, text="Diagnósticos:").grid(row=1, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(cg, textvariable=self.diagnosticos_var, width=400).grid(
            row=1, column=1, columnspan=3, sticky="ew", padx=10, pady=5
        )

        # Diagnósticos de Recetas, Labs
        ctk.CTkLabel(cg, text="Diagnósticos (Recetas/Labs):").grid(
            row=2, column=0, sticky="w", padx=10, pady=5
        )
        ctk.CTkEntry(cg, textvariable=self.diag_recetas_labs_var, width=400).grid(
            row=2, column=1, columnspan=3, sticky="ew", padx=10, pady=5
        )

        # CIE 10
        ctk.CTkLabel(cg, text="CIE-10:").grid(row=3, column=0, sticky="w", padx=10, pady=5)
        ctk.CTkEntry(cg, textvariable=self.cie10_var, width=300).grid(
            row=3, column=1, columnspan=3, sticky="ew", padx=10, pady=5
        )

        cg.grid_columnconfigure(1, weight=1)

        # ─── NOTAS ADICIONALES ────────────────────────────────────────
        nf = ctk.CTkFrame(main_frame)
        nf.pack(fill="x", pady=(0, 10))
        ctk.CTkLabel(
            nf, text="NOTAS ADICIONALES", font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=10)
        ctk.CTkLabel(nf, text="Observaciones Clínicas:").pack(anchor="w", padx=20)
        self.observaciones_text = ctk.CTkTextbox(nf, height=100)
        self.observaciones_text.pack(fill="x", padx=20, pady=(5, 10))
        ctk.CTkLabel(nf, text="Indicaciones:").pack(anchor="w", padx=20)
        self.indicaciones_text = ctk.CTkTextbox(nf, height=100)
        self.indicaciones_text.pack(fill="x", padx=20, pady=(5, 20))

        # ─── SELECCIÓN DE PLANTILLAS ────────────────────────────────────
        pf = ctk.CTkFrame(main_frame)
        pf.pack(fill="x", pady=(0, 10))
        ctk.CTkLabel(
            pf, text="SELECCIÓN DE PLANTILLAS", font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=10)
        ctk.CTkButton(
            pf, text="🔄 Actualizar Lista",
            command=self._actualizar_plantillas, width=150, height=30
        ).pack(pady=(0, 10))

        self.plantillas_scroll_frame = ctk.CTkFrame(pf, height=300)
        self.plantillas_scroll_frame.pack(fill="x", padx=20, pady=(0, 20))

        self.plantillas_vars = {}
        self._actualizar_plantillas()

        # ─── BOTONES FINALES ──────────────────────────────────────────
        bf = ctk.CTkFrame(main_frame)
        bf.pack(fill="x", pady=10)
        ctk.CTkButton(
            bf, text="GENERAR DOCUMENTOS",
            command=self._generar_documentos, width=200, height=40
        ).pack(side="left", padx=20, pady=20)
        ctk.CTkButton(
            bf, text="LIMPIAR CAMPOS",
            command=self._limpiar_campos, width=150, height=40
        ).pack(side="left", padx=10, pady=20)
        ctk.CTkButton(
            bf, text="VER HISTORIAL",
            command=self._ver_historial, width=150, height=40
        ).pack(side="right", padx=20, pady=20)

    def _configurar_eventos(self):
        """Configura los eventos de actualización automática."""
        self.fecha_nacimiento_var.trace("w", self._actualizar_edad)
        self.fecha_nacimiento_var.trace("w", self._actualizar_num_registro)

    def _actualizar_edad(self, *args):
        """Calcula edad a partir de la fecha de nacimiento."""
        try:
            fs = self.fecha_nacimiento_var.get().strip()
            if fs and '/' in fs:
                d, m, y = fs.split('/')
                dia, mes = int(d), int(m)
                año = int(y)
                if año < 100:
                    año += 2000 if año < 50 else 1900
                nac = date(año, mes, dia)
                hoy = date.today()
                edad = hoy.year - nac.year - ((hoy.month, hoy.day) < (nac.month, nac.day))
                self.edad_var.set(str(edad))
                return
            self.edad_var.set("")
        except:
            self.edad_var.set("")

    def _actualizar_num_registro(self, *args):
        """Genera el número de registro desde la fecha de nacimiento (formato: DDMMAA)."""
        try:
            fs = self.fecha_nacimiento_var.get().strip()
            if fs and '/' in fs:
                d, m, y = fs.split('/')
                dia = d.zfill(2)
                mes = m.zfill(2)
                año = y[-2:] if len(y) >= 2 else y.zfill(2)
                num_registro = f"{dia}{mes}{año}"
                self.num_registro_var.set(num_registro)
                return
            self.num_registro_var.set("")
        except:
            self.num_registro_var.set("")

    def _replace_placeholders(self, doc, data):
        """Reemplaza marcadores en .docx con Arial 10."""
        def set_arial_10(r):
            r.font.name = 'Arial'
            r.font.size = Pt(10)

        # Párrafos
        for p in doc.paragraphs:
            for key, val in data.items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(val))
                            set_arial_10(run)
        # Tablas
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, val in data.items():
                            if key in p.text:
                                for run in p.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, str(val))
                                        set_arial_10(run)
        # Headers / Footers
        for sec in doc.sections:
            for hf in (sec.header, sec.footer):
                for p in hf.paragraphs:
                    for key, val in data.items():
                        if key in p.text:
                            for run in p.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(val))
                                    set_arial_10(run)

    def _generar_documentos(self):
        """Genera y guarda los documentos seleccionados."""
        try:
            # Validaciones
            if not self.nombres_var.get().strip():
                messagebox.showerror("Error", "Nombres obligatorios")
                return
            if not self.apellido_paterno_var.get().strip():
                messagebox.showerror("Error", "Apellido paterno obligatorio")
                return
            if not self.diagnosticos_var.get().strip():
                messagebox.showerror("Error", "Diagnóstico obligatorio")
                return

            # Plantillas seleccionadas
            seleccionadas = [v for v in self.plantillas_vars.values() if v['var'].get()]
            if not seleccionadas:
                messagebox.showerror("Error", "Seleccione al menos una plantilla")
                return

            # Datos para reemplazo
            nombre_comp = f"{self.nombres_var.get().strip()} {self.apellido_paterno_var.get().strip()} {self.apellido_materno_var.get().strip()}".strip()
            data = {
                '{{NOMBRE_COMPLETO}}': nombre_comp,
                '{{NOMBRES}}': self.nombres_var.get().strip(),
                '{{APELLIDO_PATERNO}}': self.apellido_paterno_var.get().strip(),
                '{{APELLIDO_MATERNO}}': self.apellido_materno_var.get().strip(),
                '{{FECHA_NACIMIENTO}}': self.fecha_nacimiento_var.get(),
                '{{EDAD}}': self.edad_var.get(),
                '{{FECHA_INTERNACION}}': self.fecha_internacion_entry.get(),
                '{{NUM_HISTORIA}}': self.num_historia_var.get(),
                '{{NUM_REGISTRO}}': self.num_registro_var.get(),
                '{{OCUPACION}}': self.ocupacion_var.get(),
                '{{ESTADO_CIVIL}}': self.estado_civil_var.get(),
                '{{REFERENCIA1}}': self.referencia1_var.get(),
                '{{REFERENCIA2}}': self.referencia2_var.get(),
                '{{DIAGNOSTICOS}}': self.diagnosticos_var.get(),
                '{{DIAG_RECETAS_LABS}}': self.diag_recetas_labs_var.get(),
                '{{CIE10}}': self.cie10_var.get(),
                '{{SERVICIO}}': self.servicio_var.get(),
                '{{OBSERVACIONES}}': self.observaciones_text.get("1.0", "end-1c"),
                '{{INDICACIONES}}': self.indicaciones_text.get("1.0", "end-1c")
            }

            # Carpeta paciente
            nombre_carpeta = f"{nombre_comp} - {self.num_historia_var.get()} - {self.diagnosticos_var.get().strip()}"
            nombre_carpeta = re.sub(r'[<>:"/\\|?*]', '_', nombre_carpeta)
            ruta_carpeta = os.path.join(self.directorio_base, nombre_carpeta)
            os.makedirs(ruta_carpeta, exist_ok=True)

            generados, errores = [], []

            # Procesar cada plantilla
            for v in seleccionadas:
                ruta_plantilla = os.path.join(self.plantillas_dir, v['carpeta'], v['archivo'])
                if os.path.exists(ruta_plantilla):
                    try:
                        doc = Document(ruta_plantilla)
                        self._replace_placeholders(doc, data)
                        base = v['archivo'].replace('.docx', '')
                        fname = f"{base} - {nombre_comp}.docx"
                        salida = os.path.join(ruta_carpeta, fname)
                        doc.save(salida)
                        generados.append({'cat': v['carpeta'], 'file': fname})
                    except Exception as e:
                        errores.append(f"{v['archivo']}: {e}")
                else:
                    errores.append(f"No encontrado: {v['archivo']}")

            # Mensaje final
            if generados:
                self._guardar_en_excel(ruta_carpeta)
                msg = f"✅ Se generaron {len(generados)} documentos:\n"
                por_cat = {}
                for d in generados:
                    por_cat.setdefault(d['cat'], []).append(d['file'])
                for cat, files in por_cat.items():
                    msg += f"\n📁 {cat}\n"
                    for f in files:
                        msg += f" • {f}\n"
                msg += f"\n📂 Carpeta: {nombre_carpeta}"
                if errores:
                    msg += "\n\n⚠️ Errores:\n" + "\n".join(errores)
                messagebox.showinfo("Documentos Generados", msg)

                if messagebox.askyesno("Abrir Carpeta", "¿Desea abrir la carpeta?"):
                    try:
                        os.startfile(ruta_carpeta)
                    except:
                        pass
                if messagebox.askyesno("Limpiar Campos", "¿Limpiar para nuevo paciente?"):
                    self._limpiar_campos()
            else:
                msg = "❌ No se generaron documentos."
                if errores:
                    msg += "\n" + "\n".join(errores)
                messagebox.showerror("Error", msg)

        except Exception as e:
            messagebox.showerror("Error", str(e))

    def _guardar_en_excel(self, ruta_carpeta):
        """Guarda la entrada en el historial de pacientes."""
        try:
            df = pd.read_excel(self.excel_file) if os.path.exists(self.excel_file) else pd.DataFrame()
            nombre_comp = f"{self.nombres_var.get().strip()} {self.apellido_paterno_var.get().strip()} {self.apellido_materno_var.get().strip()}".strip()
            fila = {
                'Fecha_Registro': datetime.now().strftime("%d/%m/%Y %H:%M"),
                'Nombre_Completo': nombre_comp,
                'Num_Historia': self.num_historia_var.get(),
                'Num_Registro': self.num_registro_var.get(),
                'Edad': self.edad_var.get(),
                'Servicio': self.servicio_var.get(),
                'Diagnostico': self.diagnosticos_var.get(),
                'Fecha_Internacion': self.fecha_internacion_entry.get(),
                'Ruta_Carpeta': ruta_carpeta
            }
            df = pd.concat([df, pd.DataFrame([fila])], ignore_index=True)
            df.to_excel(self.excel_file, index=False)
        except Exception as e:
            messagebox.showwarning("Advertencia", f"Error guardando historial: {e}")

    def _actualizar_plantillas(self):
        """Refresca la lista de plantillas disponibles."""
        for w in self.plantillas_scroll_frame.winfo_children():
            w.destroy()
        self.plantillas_vars.clear()

        categorias = [
            ("📋 Consulta", "Consulta"),
            ("📨 Interconsulta", "Interconsulta"),
            ("💊 Recetas", "Recetas"),
            ("🧪 Laboratorios", "Laboratorios"),
            ("🔬 Procedimientos", "Procedimientos"),
            ("📊 Informes", "Informes")
        ]
        row = 0
        for title, folder in categorias:
            path = os.path.join(self.plantillas_dir, folder)
            if os.path.isdir(path):
                docs = [f for f in os.listdir(path) if f.endswith('.docx') and not f.startswith('~')]
                if docs:
                    ctk.CTkLabel(
                        self.plantillas_scroll_frame,
                        text=title,
                        font=ctk.CTkFont(size=14, weight="bold")
                    ).grid(row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 5))
                    row += 1
                    cat_var = tk.BooleanVar()
                    ctk.CTkCheckBox(
                        self.plantillas_scroll_frame,
                        text="Seleccionar todos",
                        variable=cat_var,
                        command=lambda v=cat_var, f=folder: self._toggle_categoria(v, f)
                    ).grid(row=row, column=0, columnspan=2, sticky="w", padx=30, pady=(0, 5))
                    row += 1

                    for i, doc in enumerate(docs):
                        var = tk.BooleanVar()
                        key = f"{folder}:{doc}"
                        self.plantillas_vars[key] = {
                            'var': var, 'carpeta': folder, 'archivo': doc, 'categoria_var': cat_var
                        }
                        lbl = doc.replace('.docx', '')
                        ctk.CTkCheckBox(
                            self.plantillas_scroll_frame,
                            text=lbl,
                            variable=var,
                            command=self._verificar_categoria_completa
                        ).grid(
                            row=row,
                            column=i % 2,
                            sticky="w",
                            padx=50 if i % 2 == 0 else 20,
                            pady=2
                        )
                        if i % 2 == 1:
                            row += 1
                    if len(docs) % 2 == 1:
                        row += 1
                    row += 1

        self.plantillas_scroll_frame.grid_columnconfigure(0, weight=1)
        self.plantillas_scroll_frame.grid_columnconfigure(1, weight=1)

        if not self.plantillas_vars:
            ctk.CTkLabel(
                self.plantillas_scroll_frame,
                text="⚠️ No se encontraron plantillas .docx\nColoca archivos en PLANTILLAS/",
                font=ctk.CTkFont(size=12),
                text_color="orange"
            ).grid(row=0, column=0, columnspan=2, padx=20, pady=20)

    def _toggle_categoria(self, var, folder):
        """Selecciona/deselecciona todas las plantillas de una categoría."""
        for v in self.plantillas_vars.values():
            if v['carpeta'] == folder:
                v['var'].set(var.get())

    def _verificar_categoria_completa(self):
        """Marca la casilla de categoría si todas sus plantillas están seleccionadas."""
        stats = {}
        for v in self.plantillas_vars.values():
            f = v['carpeta']
            if f not in stats:
                stats[f] = {'total': 0, 'sel': 0, 'var': v['categoria_var']}
            stats[f]['total'] += 1
            if v['var'].get():
                stats[f]['sel'] += 1
        for info in stats.values():
            info['var'].set(info['sel'] == info['total'])

    def _limpiar_campos(self):
        """Restablece todos los campos al estado inicial."""
        self.nombres_var.set("")
        self.apellido_materno_var.set("")
        self.apellido_paterno_var.set("")
        self.fecha_nacimiento_var.set("")
        self.edad_var.set("")
        self.num_historia_var.set("")
        self.num_registro_var.set("")
        self.ocupacion_var.set("")
        self.estado_civil_var.set("")
        self.referencia1_var.set("")
        self.referencia2_var.set("")
        self.diagnosticos_var.set("")
        self.diag_recetas_labs_var.set("")
        self.cie10_var.set("")
        self.servicio_var.set("Hematología")

        hoy = date.today()
        self.fecha_internacion_entry.set_date(hoy)

        self.observaciones_text.delete("1.0", "end")
        self.indicaciones_text.delete("1.0", "end")

        estado = {k: v['var'].get() for k, v in self.plantillas_vars.items()}
        self._actualizar_plantillas()
        for k, st in estado.items():
            if k in self.plantillas_vars:
                self.plantillas_vars[k]['var'].set(st)

    def _ver_historial(self):
        """Abre el archivo Excel con el historial de pacientes."""
        try:
            if os.path.exists(self.excel_file):
                os.startfile(self.excel_file)
            else:
                messagebox.showinfo("Info", "No hay historial disponible aún")
        except Exception as e:
            messagebox.showerror("Error", f"Error abriendo historial: {e}")

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = SistemaDocumentosMedicos()
    app.run()