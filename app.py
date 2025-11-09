import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import os
from datetime import datetime, date
import re
import zipfile
import io

# --- Configuración de la página de Streamlit ---
st.set_page_config(layout="wide", page_title="Sistema de Gestión de Documentos Médicos")

# --- Funciones Auxiliares ---

@st.cache_resource # Usa st.cache_resource para que esta función se ejecute una sola vez
def _crear_estructura_directorios(directorio_base, plantillas_dir, excel_file):
    """Crea las carpetas de plantillas y el archivo Excel si no existen."""
    try:
        subdirs = [
            "Consulta", "Interconsulta", "Recetas",
            "Laboratorios", "Procedimientos", "Informes"
        ]
        for sub in subdirs:
            os.makedirs(os.path.join(plantillas_dir, sub), exist_ok=True)

        if not os.path.exists(excel_file):
            df = pd.DataFrame(columns=[
                'Fecha_Registro', 'Nombre_Completo', 'Num_Historia', 'Num_Registro',
                'Edad', 'Servicio', 'Diagnostico', 'Fecha_Internacion',
                'Ocupacion', 'Estado_Civil', 'Genero', 'Residencia', 'Procedencia',
                'Domicilio', # Nueva columna para el historial
                'Residente_La_Paz',
                'Nombre_Referencia1', 'Relacion_Referencia1', 'Telefono_Referencia1', # Nuevas columnas para el historial
                'Nombre_Referencia2', 'Relacion_Referencia2', 'Telefono_Referencia2', # Nuevas columnas para el historial
                'Ruta_Carpeta'
            ])
            df.to_excel(excel_file, index=False)
        return True
    except Exception as e:
        st.error(f"Error creando estructura de directorios: {e}")
        return False

def _actualizar_edad(fecha_nacimiento_str):
    """Calcula edad a partir de la fecha de nacimiento y actualiza el estado de la sesión."""
    try:
        if fecha_nacimiento_str and '/' in fecha_nacimiento_str:
            d, m, y = fecha_nacimiento_str.split('/')
            dia, mes = int(d), int(m)
            año = int(y)
            if año < 100: # Asume años de 2 dígitos (ej. 98 -> 1998, 05 -> 2005)
                año_actual_2_dig = date.today().year % 100
                año += 1900 if año > año_actual_2_dig else 2000
            nac = date(año, mes, dia)
            hoy = date.today()
            edad = hoy.year - nac.year - ((hoy.month, hoy.day) < (nac.month, nac.day))
            st.session_state.edad = str(edad)
            return
        st.session_state.edad = ""
    except:
        st.session_state.edad = ""

def _actualizar_num_registro(fecha_nacimiento_str):
    """Genera el número de registro (DDMMAA) y actualiza el estado de la sesión."""
    try:
        if fecha_nacimiento_str and '/' in fecha_nacimiento_str:
            d, m, y = fecha_nacimiento_str.split('/')
            dia = d.zfill(2)
            mes = m.zfill(2)
            año = y[-2:] if len(y) >= 2 else y.zfill(2) # Obtener los últimos 2 dígitos
            num_registro = f"{dia}{mes}{año}"
            st.session_state.num_registro = num_registro
            return
        st.session_state.num_registro = ""
    except:
        st.session_state.num_registro = ""

def _replace_placeholders(doc, data):
    """
    Reemplaza marcadores en .docx con los datos proporcionados.
    Intenta preservar el formato original de los runs.
    """
    # Párrafos
    for p in doc.paragraphs:
        for key, val in data.items():
            if key in p.text:
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(val))
    # Tablas
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in data.items():
                        if key in p.text:
                            for run in p.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(val))
    # Encabezados y pies de página
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                for key, val in data.items():
                    if key in p.text:
                        for run in p.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, str(val))

def _guardar_en_excel(excel_file, ruta_carpeta_simulada):
    """Guarda la entrada del paciente en el historial Excel."""
    try:
        df = pd.read_excel(excel_file) if os.path.exists(excel_file) else pd.DataFrame()
        nombre_comp = f"{st.session_state.nombres.strip()} {st.session_state.apellido_paterno.strip()} {st.session_state.apellido_materno.strip()}".strip()
        fila = {
            'Fecha_Registro': datetime.now().strftime("%d/%m/%Y %H:%M"),
            'Nombre_Completo': nombre_comp,
            'Num_Historia': st.session_state.num_historia,
            'Num_Registro': st.session_state.num_registro,
            'Edad': st.session_state.edad,
            'Servicio': st.session_state.servicio,
            'Diagnostico': st.session_state.diagnosticos,
            'Fecha_Internacion': st.session_state.fecha_internacion.strftime("%d/%m/%Y"),
            'Ocupacion': st.session_state.ocupacion,
            'Estado_Civil': st.session_state.estado_civil,
            'Genero': st.session_state.genero,
            'Residencia': st.session_state.residencia,
            'Procedencia': st.session_state.procedencia,
            'Domicilio': st.session_state.domicilio, # NUEVO
            'Residente_La_Paz': "Sí" if st.session_state.es_residente_la_paz else "No",
            'Nombre_Referencia1': st.session_state.n_referencia1, # NUEVO
            'Relacion_Referencia1': st.session_state.referencia1,
            'Telefono_Referencia1': st.session_state.telefono_referencia1, # NUEVO
            'Nombre_Referencia2': st.session_state.n_referencia2, # NUEVO
            'Relacion_Referencia2': st.session_state.referencia2,
            'Telefono_Referencia2': st.session_state.telefono_referencia2, # NUEVO
            'Ruta_Carpeta': ruta_carpeta_simulada
        }
        df = pd.concat([df, pd.DataFrame([fila])], ignore_index=True)
        df.to_excel(excel_file, index=False)
        return True
    except Exception as e:
        st.warning(f"Advertencia: Error guardando historial en Excel: {e}")
        return False

def _limpiar_campos():
    """Restablece todos los campos de entrada a su estado inicial."""
    st.session_state.nombres = ""
    st.session_state.apellido_materno = ""
    st.session_state.apellido_paterno = ""
    st.session_state.fecha_nacimiento_str = ""
    st.session_state.edad = ""
    st.session_state.num_historia = ""
    st.session_state.num_registro = ""
    st.session_state.ocupacion = ""
    st.session_state.estado_civil = ""
    st.session_state.residencia = ""
    st.session_state.genero = "Masculino"
    st.session_state.procedencia = ""
    st.session_state.domicilio = "" # NUEVO
    st.session_state.es_residente_la_paz = False
    st.session_state.n_referencia1 = ""
    st.session_state.referencia1 = ""
    st.session_state.telefono_referencia1 = "" # NUEVO
    st.session_state.n_referencia2 = ""
    st.session_state.referencia2 = ""
    st.session_state.telefono_referencia2 = "" # NUEVO
    st.session_state.diagnosticos = ""
    st.session_state.diag_recetas_labs = ""
    st.session_state.cie10 = ""
    st.session_state.servicio = "Hematología"
    st.session_state.observaciones = ""
    st.session_state.indicaciones = ""
    st.session_state.fecha_internacion = date.today()

    for key in list(st.session_state.plantillas_vars.keys()):
        st.session_state.plantillas_vars[key] = False
    for key in list(st.session_state.keys()):
        if key.startswith("select_all_"):
            st.session_state[key] = False

def _generar_documentos_callback(directorio_base, plantillas_dir, excel_file):
    """
    Función callback para generar los documentos seleccionados.
    Prepara un archivo ZIP con los documentos y ofrece la descarga.
    """
    # Validaciones básicas
    if not st.session_state.nombres.strip():
        st.error("Por favor, ingrese los Nombres del paciente.")
        return
    if not st.session_state.apellido_paterno.strip():
        st.error("Por favor, ingrese el Apellido Paterno del paciente.")
        return
    if not st.session_state.diagnosticos.strip():
        st.error("Por favor, ingrese el Diagnóstico del paciente.")
        return

    # Obtener plantillas seleccionadas
    seleccionadas = []
    for k, v in st.session_state.plantillas_vars.items():
        if v:
            folder, filename = k.split(':', 1)
            seleccionadas.append({'carpeta': folder, 'archivo': filename})

    if not seleccionadas:
        st.error("Seleccione al menos una plantilla para generar documentos.")
        return

    # Preparar datos para reemplazo en la plantilla
    nombre_comp = f"{st.session_state.nombres.strip()} {st.session_state.apellido_paterno.strip()} {st.session_state.apellido_materno.strip()}".strip()
    data = {
        '{{NOMBRE_COMPLETO}}': nombre_comp,
        '{{NOMBRES}}': st.session_state.nombres.strip(),
        '{{APELLIDO_PATERNO}}': st.session_state.apellido_paterno.strip(),
        '{{APELLIDO_MATERNO}}': st.session_state.apellido_materno.strip(),
        '{{FECHA_NACIMIENTO}}': st.session_state.fecha_nacimiento_str,
        '{{EDAD}}': st.session_state.edad,
        '{{FECHA_INTERNACION}}': st.session_state.fecha_internacion.strftime("%d/%m/%Y"),
        '{{NUM_HISTORIA}}': st.session_state.num_historia,
        '{{NUM_REGISTRO}}': st.session_state.num_registro,
        '{{OCUPACION}}': st.session_state.ocupacion,
        '{{ESTADO_CIVIL}}': st.session_state.estado_civil,
        '{{RESIDENCIA}}': st.session_state.residencia,
        '{{GENERO}}': st.session_state.genero,
        '{{PROCEDENCIA}}': st.session_state.procedencia,
        '{{DOMICILIO}}': st.session_state.domicilio, # NUEVO MARCADO
        '{{ES_RESIDENTE_LA_PAZ}}': "Sí" if st.session_state.es_residente_la_paz else "No",
        '{{N_REFERENCIA1}}': st.session_state.n_referencia1,
        '{{REFERENCIA1}}': st.session_state.referencia1,
        '{{TELEFONO_REFERENCIA1}}': st.session_state.telefono_referencia1, # NUEVO MARCADO
        '{{N_REFERENCIA2}}': st.session_state.n_referencia2,
        '{{REFERENCIA2}}': st.session_state.referencia2,
        '{{TELEFONO_REFERENCIA2}}': st.session_state.telefono_referencia2, # NUEVO MARCADO
        '{{DIAGNOSTICOS}}': st.session_state.diagnosticos,
        '{{DIAG_RECETAS_LABS}}': st.session_state.diag_recetas_labs,
        '{{CIE10}}': st.session_state.cie10,
        '{{SERVICIO}}': st.session_state.servicio,
        '{{OBSERVACIONES}}': st.session_state.observaciones,
        '{{INDICACIONES}}': st.session_state.indicaciones
    }

    # Crear nombre de carpeta "virtual" para el ZIP y sanitizarlo
    nombre_carpeta_raw = f"{nombre_comp} - {st.session_state.num_historia} - {st.session_state.diagnosticos.strip()}"
    nombre_carpeta_sanitized = re.sub(r'[<>:"/\\|?*]', '_', nombre_carpeta_raw)

    generados_buffer = io.BytesIO()
    with zipfile.ZipFile(generados_buffer, 'w') as zf:
        generados, errores = [], []

        for v in seleccionadas:
            ruta_plantilla = os.path.join(plantillas_dir, v['carpeta'], v['archivo'])
            if os.path.exists(ruta_plantilla):
                try:
                    doc = Document(ruta_plantilla)
                    _replace_placeholders(doc, data)
                    base = v['archivo'].replace('.docx', '')
                    fname = f"{base} - {nombre_comp}.docx"

                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    zf.writestr(os.path.join(nombre_carpeta_sanitized, fname), doc_buffer.getvalue())
                    generados.append({'cat': v['carpeta'], 'file': fname})
                except Exception as e:
                    errores.append(f"{v['archivo']}: {e}")
            else:
                errores.append(f"Plantilla no encontrada: {v['archivo']}")

    if generados:
        _guardar_en_excel(excel_file, nombre_carpeta_sanitized)
        
        msg = f"✅ Se generaron {len(generados)} documentos. Haga clic en 'Descargar Documentos' para obtener el archivo ZIP.\n"
        por_cat = {}
        for d in generados:
            por_cat.setdefault(d['cat'], []).append(d['file'])
        for cat, files in por_cat.items():
            msg += f"\n📁 **{cat}**\n"
            for f in files:
                msg += f" • {f}\n"
        if errores:
            msg += "\n\n⚠️ **Errores al generar algunos documentos:**\n" + "\n".join(errores)
        
        st.success(msg)

        generados_buffer.seek(0)
        st.download_button(
            label=f"Descargar Documentos Generados ({len(generados)})",
            data=generados_buffer.getvalue(),
            file_name=f"{nombre_carpeta_sanitized}.zip",
            mime="application/zip",
            key="download_docs_button"
        )
        
        if st.button("Limpiar campos para nuevo paciente", key="clear_after_gen"):
            _limpiar_campos()
            st.rerun()
    else:
        msg = "❌ No se generaron documentos."
        if errores:
            msg += "\n" + "\n".join(errores)
        st.error(msg)

# --- Lógica principal de la aplicación Streamlit ---
def main():
    directorio_base = os.getcwd()
    plantillas_dir = os.path.join(directorio_base, "PLANTILLAS")
    excel_file = os.path.join(directorio_base, "pacientes.xlsx")

    # Inicializar variables de estado de sesión si no existen
    if 'initialized' not in st.session_state:
        st.session_state.initialized = True
        st.session_state.nombres = ""
        st.session_state.apellido_materno = ""
        st.session_state.apellido_paterno = ""
        st.session_state.fecha_nacimiento_str = ""
        st.session_state.edad = ""
        st.session_state.num_historia = ""
        st.session_state.num_registro = ""
        st.session_state.ocupacion = ""
        st.session_state.estado_civil = ""
        st.session_state.residencia = ""
        st.session_state.genero = "Masculino"
        st.session_state.procedencia = ""
        st.session_state.domicilio = "" # NUEVO
        st.session_state.es_residente_la_paz = False
        st.session_state.n_referencia1 = ""
        st.session_state.referencia1 = ""
        st.session_state.telefono_referencia1 = "" # NUEVO
        st.session_state.n_referencia2 = ""
        st.session_state.referencia2 = ""
        st.session_state.telefono_referencia2 = "" # NUEVO
        st.session_state.diagnosticos = ""
        st.session_state.diag_recetas_labs = ""
        st.session_state.cie10 = ""
        st.session_state.servicio = "Hematología"
        st.session_state.observaciones = ""
        st.session_state.indicaciones = ""
        st.session_state.fecha_internacion = date.today()
        st.session_state.plantillas_vars = {}

    if not _crear_estructura_directorios(directorio_base, plantillas_dir, excel_file):
        st.stop()

    st.title("Sistema de Gestión de Documentos Médicos")

    # --- Sección de DATOS DEL PACIENTE ---
    st.header("1. Datos del Paciente")
    with st.container(border=True):
        col1, col2 = st.columns(2)
        with col1:
            st.text_input("Nombres:", key="nombres")
            st.text_input("Apellido Paterno:", key="apellido_paterno")
            st.text_input("Fecha Nacimiento (DD/MM/YY):", key="fecha_nacimiento_str",
                           on_change=lambda: (_actualizar_edad(st.session_state.fecha_nacimiento_str),
                                              _actualizar_num_registro(st.session_state.fecha_nacimiento_str)),
                           placeholder="DD/MM/YY", help="Formato: DD/MM/AA o DD/MM/AAAA")
            st.date_input("Fecha Internación:", key="fecha_internacion", value=st.session_state.fecha_internacion, format="DD/MM/YYYY")
            st.text_input("N° Historia Clínica:", key="num_historia")
            st.text_input("Ocupación:", key="ocupacion")
            st.text_input("Estado Civil:", key="estado_civil")
            st.text_input("Residencia:", key="residencia", help="Localidad o ciudad de residencia (ej. La Paz, El Alto)")
            st.text_input("Domicilio:", key="domicilio", help="Dirección completa del domicilio del paciente (calle, número, zona).") # NUEVO
            st.text_input("Nombre Referencia 1:", key="n_referencia1", help="Nombre completo de la primera persona de referencia")
            st.text_input("Referencia 1 (Parentesco/Relación):", key="referencia1", help="Parentesco o relación de la referencia 1 con el paciente")
            st.text_input("Teléfono Referencia 1:", key="telefono_referencia1", help="Número de teléfono de la referencia 1") # NUEVO

        with col2:
            st.text_input("Apellido Materno:", key="apellido_materno")
            st.text_input("Edad:", key="edad", disabled=True, help="Calculada automáticamente")
            st.text_input("N° Registro:", key="num_registro", disabled=True, help="Calculado automáticamente (DDMMAA)")
            st.selectbox("Género:", options=["Masculino", "Femenino", "Otro"], key="genero", help="Género del paciente")
            st.text_input("Procedencia:", key="procedencia", help="Lugar de donde proviene el paciente (si es diferente a residencia)")
            st.checkbox("Residente en La Paz:", key="es_residente_la_paz", help="Marque si el paciente reside en La Paz")
            st.text_input("Nombre Referencia 2:", key="n_referencia2", help="Nombre completo de la segunda persona de referencia")
            st.text_input("Referencia 2 (Parentesco/Relación):", key="referencia2", help="Parentesco o relación de la referencia 2 con el paciente")
            st.text_input("Teléfono Referencia 2:", key="telefono_referencia2", help="Número de teléfono de la referencia 2") # NUEVO


    # --- Sección de DATOS CLÍNICOS ---
    st.header("2. Datos Clínicos")
    with st.container(border=True):
        col1, col2 = st.columns(2)
        with col1:
            servicios = ["Hematología", "Medicina Interna", "Oncología Clínica", "Oncología Quirúrgica"]
            st.selectbox("Servicio:", options=servicios, key="servicio")
            st.text_input("Diagnósticos:", key="diagnosticos", help="Diagnósticos principales del paciente")
        with col2:
            st.text_input("Diagnósticos (para Recetas/Labs):", key="diag_recetas_labs", help="Diagnósticos específicos para recetas o laboratorios")
            st.text_input("CIE-10:", key="cie10", help="Código de la Clasificación Internacional de Enfermedades (CIE-10)")

    # --- Sección de NOTAS ADICIONALES ---
    st.header("3. Notas Adicionales")
    with st.container(border=True):
        st.text_area("Observaciones Clínicas:", key="observaciones", height=100, help="Notas adicionales relevantes para el caso clínico.")
        st.text_area("Indicaciones:", key="indicaciones", height=100, help="Instrucciones médicas o terapéuticas.")

    # --- Sección de SELECCIÓN DE PLANTILLAS ---
    st.header("4. Selección de Plantillas")
    
    def _render_plantillas_selection():
        categorias = [
            ("📋 Consulta", "Consulta"),
            ("📨 Interconsulta", "Interconsulta"),
            ("💊 Recetas", "Recetas"),
            ("🧪 Laboratorios", "Laboratorios"),
            ("🔬 Procedimientos", "Procedimientos"),
            ("📊 Informes", "Informes")
        ]
        
        st.button("🔄 Actualizar Lista de Plantillas", key="refresh_templates", help="Vuelve a escanear la carpeta PLANTILLAS para nuevas plantillas.")

        template_found = False
        with st.container(border=True):
            for title, folder in categorias:
                path = os.path.join(plantillas_dir, folder)
                if os.path.isdir(path):
                    docs = [f for f in os.listdir(path) if f.endswith('.docx') and not f.startswith('~')]
                    if docs:
                        template_found = True
                        st.subheader(title)
                        
                        select_all_key = f"select_all_{folder}"
                        
                        all_selected_in_category = True
                        for doc in docs:
                            key = f"{folder}:{doc}"
                            if key not in st.session_state.plantillas_vars or not st.session_state.plantillas_vars[key]:
                                all_selected_in_category = False
                                break
                        
                        if select_all_key not in st.session_state:
                            st.session_state[select_all_key] = all_selected_in_category
                        else:
                            if st.session_state[select_all_key] != all_selected_in_category:
                                st.session_state[select_all_key] = all_selected_in_category
                        
                        def toggle_category(category_folder):
                            for doc_name in os.listdir(os.path.join(plantillas_dir, category_folder)):
                                if doc_name.endswith('.docx') and not doc_name.startswith('~'):
                                    key = f"{category_folder}:{doc_name}"
                                    st.session_state.plantillas_vars[key] = st.session_state[f"select_all_{category_folder}"]
                            st.rerun()

                        st.checkbox("Seleccionar todos en esta categoría", key=select_all_key, on_change=toggle_category, args=(folder,), help="Marca/desmarca todas las plantillas de esta sección.")
                        
                        cols_per_row = 2
                        current_cols = st.columns(cols_per_row)
                        col_idx = 0
                        
                        for doc in docs:
                            key = f"{folder}:{doc}"
                            if key not in st.session_state.plantillas_vars:
                                st.session_state.plantillas_vars[key] = False
                            
                            lbl = doc.replace('.docx', '')
                            def check_category_all_status(current_folder):
                                all_sub_selected = True
                                for sub_doc in os.listdir(os.path.join(plantillas_dir, current_folder)):
                                    if sub_doc.endswith('.docx') and not sub_doc.startswith('~'):
                                        sub_key = f"{current_folder}:{sub_doc}"
                                        if not st.session_state.plantillas_vars.get(sub_key, False):
                                            all_sub_selected = False
                                            break
                                st.session_state[f"select_all_{current_folder}"] = all_sub_selected
                                st.rerun()
                            
                            current_cols[col_idx].checkbox(lbl, key=key, on_change=check_category_all_status, args=(folder,))
                            col_idx = (col_idx + 1) % cols_per_row
                            if col_idx == 0:
                                current_cols = st.columns(cols_per_row)

            if not template_found:
                st.warning("⚠️ No se encontraron plantillas .docx. Asegúrate de colocar archivos en la carpeta `PLANTILLAS/` con subcarpetas para categorías (ej. `PLANTILLAS/Consulta/`).")

    _render_plantillas_selection()


    # --- Sección de ACCIONES FINALES ---
    st.header("5. Acciones")
    col_gen, col_clear, col_hist = st.columns(3)

    with col_gen:
        st.button(
            "GENERAR DOCUMENTOS",
            on_click=_generar_documentos_callback,
            args=(directorio_base, plantillas_dir, excel_file),
            type="primary",
            use_container_width=True
        )
    with col_clear:
        st.button("LIMPIAR CAMPOS", on_click=_limpiar_campos, use_container_width=True)
    with col_hist:
        def _toggle_historial_visibility():
            if not st.session_state.get('show_historial', False):
                try:
                    if os.path.exists(excel_file):
                        df_historial = pd.read_excel(excel_file)
                        st.session_state.historial_data = df_historial
                    else:
                        st.session_state.historial_data = pd.DataFrame()
                except Exception as e:
                    st.error(f"Error cargando historial: {e}")
                    st.session_state.historial_data = pd.DataFrame()
            st.session_state.show_historial = not st.session_state.get('show_historial', False)

        if 'show_historial' not in st.session_state:
            st.session_state.show_historial = False
        if 'historial_data' not in st.session_state:
            st.session_state.historial_data = pd.DataFrame()

        st.button("VER HISTORIAL", on_click=_toggle_historial_visibility, use_container_width=True)

    if st.session_state.get('show_historial'):
        st.subheader("Historial de Pacientes")
        if not st.session_state.historial_data.empty:
            st.dataframe(st.session_state.historial_data, use_container_width=True)
            
            excel_buffer = io.BytesIO()
            st.session_state.historial_data.to_excel(excel_buffer, index=False)
            excel_buffer.seek(0)
            st.download_button(
                label="Descargar Historial (Excel)",
                data=excel_buffer.getvalue(),
                file_name="pacientes_historial.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_historial_button"
            )
        else:
            st.info("El historial está vacío. Genere documentos para empezar a registrar pacientes.")

if __name__ == "__main__":
    main()
